import argparse
from openai import OpenAI
import tempfile
from docx import Document
import os

def transcribe_audio(file_path, api_key):
    client = OpenAI(api_key=api_key)
    with open(file_path, "rb") as audio_file:
        transcript = client.audio.transcriptions.create(
            model="whisper-1", 
            file=audio_file
        )
    return transcript.text

def save_transcript_to_word(transcript, output_file):
    doc = Document()
    doc.add_heading("Transcription", level=1)
    doc.add_paragraph(transcript)
    doc.save(output_file)

def main():
    parser = argparse.ArgumentParser(description="Transcribe audio file and save as Word document")
    parser.add_argument("--input_file", help="Path to the input audio file")
    parser.add_argument("--output_file", help="Path to the output Word document")
    parser.add_argument("--api_key", help="OpenAI API Key")
    args = parser.parse_args()

    if not args.api_key:
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OpenAI API Key must be provided either as an argument or as an environment variable OPENAI_API_KEY")
    else:
        api_key = args.api_key

    print("Transcribing...")
    try:
        transcript = transcribe_audio(args.input_file, api_key)
        print("Transcription complete. Saving to Word document...")
        save_transcript_to_word(transcript, args.output_file)
        print(f"Transcription saved to {args.output_file}")
    except Exception as e:
        print(f"An error occurred during transcription: {e}")

if __name__ == "__main__":
    main()