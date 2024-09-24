import streamlit as st
from openai import OpenAI
import tempfile
from docx import Document
from io import BytesIO

# Streamlit App
st.title("Transcriber")

st.write("Upload an audio file and transcribe it using OpenAI Whisper API.")

# Sidebar input for OpenAI API key
st.sidebar.title("API Configuration")
openai_api_key = st.sidebar.text_input("Enter your OpenAI API Key", type="password")

# Add a save button for the API key
if st.sidebar.button("Save API Key"):
    if openai_api_key:
        st.session_state['openai_api_key'] = openai_api_key
        st.sidebar.success("API Key saved successfully!")
    else:
        st.sidebar.error("Please enter an API Key before saving.")

# Check if the API key is saved in the session state
if 'openai_api_key' in st.session_state:
    openai_api_key = st.session_state['openai_api_key']
    st.sidebar.success("API Key is set.")
else:
    st.sidebar.warning("Please enter your OpenAI API Key and click 'Save API Key' to proceed.")

# File uploader
uploaded_file = st.file_uploader("Choose an audio file", type=["mp3", "wav", "m4a"])

# Function to transcribe the audio using the updated API
def transcribe_audio(file_path):
    client = OpenAI(api_key=openai_api_key)
    with open(file_path, "rb") as audio_file:
        transcript = client.audio.transcriptions.create(
            model="whisper-1", 
            file=audio_file
        )
    return transcript.text

# Function to save transcript as Word document
def save_transcript_to_word(transcript):
    doc = Document()
    doc.add_heading("Transcription", level=1)
    doc.add_paragraph(transcript)

    # Save to a BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# Transcribe button
if uploaded_file is not None:
    if 'openai_api_key' in st.session_state and st.button("Transcribe"):
        # Save the uploaded file to a temporary location
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.getbuffer())
            temp_file_path = temp_file.name

        st.write("Transcribing...")

        try:
            # Transcribe the audio
            transcript = transcribe_audio(temp_file_path)

            # Display the transcription in Markdown format
            st.subheader("Transcription")
            st.markdown(transcript)

            # Save the transcription to a Word document
            word_doc = save_transcript_to_word(transcript)

            # Provide download link for the Word document
            st.download_button(
                label="Download Transcription as Word",
                data=word_doc,
                file_name="transcription.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"An error occurred during transcription: {e}")
    elif 'openai_api_key' not in st.session_state:
        st.warning("Please save your API Key in the sidebar before transcribing.")
